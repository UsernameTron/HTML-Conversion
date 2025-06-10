"""Advanced DOCX text extraction with formatting preservation."""

import logging
import base64
from io import BytesIO
from typing import Tuple, Optional, Dict, Any
import streamlit as st

# DOCX processing libraries
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Advanced DOCX processor with formatting preservation."""
    
    def __init__(self):
        """Initialize DOCX processor."""
        self.available = PYTHON_DOCX_AVAILABLE
        
        if not self.available:
            logger.error("python-docx library not available")
    
    def extract_text_from_docx(self, file_obj, filename: str) -> Tuple[bool, str]:
        """
        Extract text from DOCX with formatting preservation.
        
        Args:
            file_obj: File object containing DOCX data
            filename: Name of the DOCX file
            
        Returns:
            Tuple of (success, extracted_content_or_error)
        """
        if not self.available:
            return False, "python-docx library not available. Please install python-docx."
        
        # Reset file pointer
        file_obj.seek(0)
        file_data = file_obj.read()
        file_size = len(file_data)
        
        # Validate file size (50MB limit)
        if file_size > 50 * 1024 * 1024:
            return False, "DOCX file too large (max 50MB)"
        
        logger.info(f"Processing DOCX: {filename} ({file_size:,} bytes)")
        
        try:
            # Create BytesIO stream for docx processing
            docx_stream = BytesIO(file_data)
            
            # Load document
            doc = Document(docx_stream)
            
            # Extract content with formatting
            html_content = self._extract_with_formatting(doc, filename)
            
            if html_content.strip():
                logger.info(f"Successfully extracted {len(html_content):,} characters from DOCX")
                return True, self._format_docx_content(html_content, filename, file_size)
            else:
                return True, self._create_empty_placeholder(filename, file_size)
                
        except Exception as e:
            logger.warning(f"DOCX extraction failed for {filename}: {str(e)}")
            return True, self._create_error_placeholder(filename, file_size, str(e))
    
    def _extract_with_formatting(self, doc: Document, filename: str) -> str:
        """Extract content while preserving formatting."""
        html_parts = []
        
        try:
            # Extract document properties
            core_props = doc.core_properties
            if core_props.title:
                html_parts.append(f'<h1>{core_props.title}</h1>')
            
            # Process paragraphs
            for para in doc.paragraphs:
                para_html = self._process_paragraph(para)
                if para_html:
                    html_parts.append(para_html)
            
            # Process tables
            for table_num, table in enumerate(doc.tables, 1):
                table_html = self._process_table(table, table_num)
                if table_html:
                    html_parts.append(table_html)
            
            return '\\n'.join(html_parts)
            
        except Exception as e:
            logger.warning(f"Error extracting DOCX content: {str(e)}")
            return ""
    
    def _process_paragraph(self, paragraph) -> str:
        """Process a paragraph with formatting."""
        if not paragraph.text.strip():
            return ""
        
        try:
            # Determine paragraph style
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            
            # Handle headings
            if 'heading' in style_name:
                level = self._extract_heading_level(style_name)
                return f'<h{level}>{self._process_runs(paragraph.runs)}</h{level}>'
            
            # Handle list items
            elif 'list' in style_name or paragraph.text.strip().startswith(('•', '-', '*')):
                return f'<li>{self._process_runs(paragraph.runs)}</li>'
            
            # Handle quotes
            elif 'quote' in style_name:
                return f'<blockquote>{self._process_runs(paragraph.runs)}</blockquote>'
            
            # Regular paragraph
            else:
                content = self._process_runs(paragraph.runs)
                if content.strip():
                    # Apply paragraph alignment if specified
                    align_style = self._get_alignment_style(paragraph)
                    if align_style:
                        return f'<p style="{align_style}">{content}</p>'
                    else:
                        return f'<p>{content}</p>'
        
        except Exception as e:
            logger.warning(f"Error processing paragraph: {str(e)}")
            return f'<p>{paragraph.text}</p>'
        
        return ""
    
    def _process_runs(self, runs) -> str:
        """Process text runs with formatting."""
        if not runs:
            return ""
        
        processed_text = []
        
        for run in runs:
            if not run.text:
                continue
            
            text = run.text
            
            try:
                # Apply formatting
                if run.bold:
                    text = f'<strong>{text}</strong>'
                
                if run.italic:
                    text = f'<em>{text}</em>'
                
                if run.underline:
                    text = f'<u>{text}</u>'
                
                # Handle font color (if available)
                if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
                    color = self._rgb_to_hex(run.font.color.rgb)
                    text = f'<span style="color: {color};">{text}</span>'
                
                # Handle font size (if available)
                if hasattr(run.font, 'size') and run.font.size:
                    size_pt = run.font.size.pt
                    text = f'<span style="font-size: {size_pt}pt;">{text}</span>'
                
            except Exception as e:
                logger.debug(f"Error applying run formatting: {str(e)}")
                # Use plain text if formatting fails
                text = run.text
            
            processed_text.append(text)
        
        return ''.join(processed_text)
    
    def _process_table(self, table, table_num: int) -> str:
        """Process a table into HTML."""
        try:
            html_parts = [
                f'<div style="margin: 20px 0;">',
                f'<h4>Table {table_num}</h4>',
                '<table style="border-collapse: collapse; width: 100%; border: 1px solid #ddd;">',
            ]
            
            for row_num, row in enumerate(table.rows):
                html_parts.append('<tr>')
                
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    
                    # Use header styling for first row
                    if row_num == 0:
                        html_parts.append(
                            f'<th style="border: 1px solid #ddd; padding: 8px; background: #f5f5f5; font-weight: bold;">'
                            f'{cell_text}</th>'
                        )
                    else:
                        html_parts.append(
                            f'<td style="border: 1px solid #ddd; padding: 8px;">{cell_text}</td>'
                        )
                
                html_parts.append('</tr>')
            
            html_parts.extend(['</table>', '</div>'])
            
            return '\\n'.join(html_parts)
            
        except Exception as e:
            logger.warning(f"Error processing table {table_num}: {str(e)}")
            return f'<p><em>Table {table_num} - processing error</em></p>'
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            # Look for numbers in style name
            import re
            numbers = re.findall(r'\\d+', style_name)
            if numbers:
                level = int(numbers[0])
                return min(max(level, 1), 6)  # Clamp between 1-6
        except:
            pass
        
        return 2  # Default to h2
    
    def _get_alignment_style(self, paragraph) -> str:
        """Get CSS alignment style for paragraph."""
        try:
            if hasattr(paragraph, 'alignment') and paragraph.alignment:
                alignment_map = {
                    WD_ALIGN_PARAGRAPH.CENTER: "text-align: center;",
                    WD_ALIGN_PARAGRAPH.RIGHT: "text-align: right;",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "text-align: justify;",
                }
                return alignment_map.get(paragraph.alignment, "")
        except:
            pass
        
        return ""
    
    def _rgb_to_hex(self, rgb) -> str:
        """Convert RGB to hex color."""
        try:
            return f"#{rgb:06x}"
        except:
            return "#000000"
    
    def _format_docx_content(self, content: str, filename: str, file_size: int) -> str:
        """Format extracted DOCX content with metadata."""
        file_size_mb = file_size / (1024 * 1024)
        
        header = f'''<div style="border: 2px solid #4caf50; border-radius: 12px; padding: 20px; margin: 20px 0; 
                     background: linear-gradient(135deg, #e8f5e8 0%, #c8e6c9 100%);">
    <div style="display: flex; align-items: center; margin-bottom: 15px;">
        <div style="font-size: 32px; margin-right: 15px;">📝</div>
        <div>
            <h2 style="color: #2e7d32; margin: 0; font-size: 20px;">Word Document Content</h2>
            <h3 style="color: #2d3748; margin: 5px 0 0 0; font-size: 16px;">{filename}</h3>
        </div>
    </div>
    <div style="background: white; padding: 12px; border-radius: 6px; border-left: 4px solid #4caf50;">
        <p style="margin: 0; font-size: 13px; color: #4a5568;"><strong>File Size:</strong> {file_size_mb:.1f} MB</p>
        <p style="margin: 4px 0 0 0; font-size: 13px; color: #4a5568;"><strong>Extraction:</strong> Successful with formatting preservation</p>
        <p style="margin: 4px 0 0 0; font-size: 13px; color: #4a5568;"><strong>Content Length:</strong> {len(content):,} characters</p>
    </div>
</div>

<div style="margin: 20px 0;">
{content}
</div>'''
        
        return header
    
    def _create_empty_placeholder(self, filename: str, file_size: int) -> str:
        """Create placeholder for empty documents."""
        file_size_mb = file_size / (1024 * 1024)
        
        return f'''<div style="border: 2px solid #9e9e9e; border-radius: 12px; padding: 30px; margin: 20px 0; 
                   background: linear-gradient(135deg, #f5f5f5 0%, #eeeeee 100%); text-align: center;">
    <div style="font-size: 48px; margin-bottom: 15px;">📄</div>
    <h2 style="color: #616161; margin-bottom: 10px; font-size: 24px;">Empty Document</h2>
    <h3 style="color: #2d3748; margin-bottom: 15px; font-size: 18px;">{filename}</h3>
    <div style="background: white; padding: 15px; border-radius: 8px; margin: 15px 0; border-left: 4px solid #9e9e9e;">
        <p style="margin: 0; font-size: 14px; color: #4a5568;"><strong>File Size:</strong> {file_size_mb:.1f} MB</p>
        <p style="margin: 5px 0 0 0; font-size: 14px; color: #4a5568;"><strong>Status:</strong> Document processed successfully</p>
        <p style="margin: 5px 0 0 0; font-size: 14px; color: #4a5568;"><strong>Content:</strong> No text content found</p>
    </div>
    <p style="color: #757575; font-size: 14px; line-height: 1.5; margin-bottom: 0;">
        This document appears to be empty or contains only non-text elements.
    </p>
</div>'''
    
    def _create_error_placeholder(self, filename: str, file_size: int, error_message: str) -> str:
        """Create placeholder when extraction fails."""
        file_size_mb = file_size / (1024 * 1024)
        
        return f'''<div style="border: 2px solid #ff9800; border-radius: 12px; padding: 30px; margin: 20px 0; 
                   background: linear-gradient(135deg, #fff3e0 0%, #ffe0b2 100%); text-align: center;">
    <div style="font-size: 48px; margin-bottom: 15px;">📝</div>
    <h2 style="color: #ef6c00; margin-bottom: 10px; font-size: 24px;">Document Processing Status</h2>
    <h3 style="color: #2d3748; margin-bottom: 15px; font-size: 18px;">{filename}</h3>
    <div style="background: white; padding: 15px; border-radius: 8px; margin: 15px 0; border-left: 4px solid #ff9800;">
        <p style="margin: 0; font-size: 14px; color: #4a5568;"><strong>File Size:</strong> {file_size_mb:.1f} MB</p>
        <p style="margin: 5px 0 0 0; font-size: 14px; color: #4a5568;"><strong>Status:</strong> Processing attempted</p>
        <p style="margin: 5px 0 0 0; font-size: 14px; color: #4a5568;"><strong>Issue:</strong> {error_message}</p>
    </div>
    <div style="background: #fff3e0; padding: 15px; border-radius: 8px; border: 1px solid #ffcc02;">
        <p style="color: #ef6c00; font-size: 14px; line-height: 1.5; margin: 0;">
            <strong>📋 Troubleshooting:</strong><br>
            • Check if the document is corrupted<br>
            • Try opening in Microsoft Word first<br>
            • Save as a newer .docx format<br>
            • Copy and paste content manually for formatting
        </p>
    </div>
</div>'''